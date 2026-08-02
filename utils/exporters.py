from __future__ import annotations

from functools import lru_cache
from io import BytesIO
import json
from pathlib import Path
import platform
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont, features

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError:  # Native Pillow/Raqm is preferred; these are fallback-only.
    arabic_reshaper = None
    get_display = None


PAGE_WIDTH = 1654
PAGE_HEIGHT = 2339
MARGIN = 110
CONTENT_WIDTH = PAGE_WIDTH - (MARGIN * 2)
CARD_PADDING = 38

BACKGROUND = (247, 249, 252)
CARD_BACKGROUND = (255, 255, 255)
TITLE_COLOR = (15, 23, 42)
TEXT_COLOR = (30, 41, 59)
MUTED_COLOR = (100, 116, 139)
ACCENT_COLOR = (91, 70, 229)
BORDER_COLOR = (218, 224, 234)
HEADER_BACKGROUND = (12, 18, 38)

ARABIC_PATTERN = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
LATIN_PATTERN = re.compile(r"[A-Za-z]")


def _font_candidates(bold: bool = False) -> list[str]:
    system = platform.system().lower()
    candidates: list[str] = []

    if "windows" in system:
        folder = Path("C:/Windows/Fonts")
        names = (
            ["tahomabd.ttf", "arialbd.ttf", "seguisb.ttf"]
            if bold
            else ["tahoma.ttf", "arial.ttf", "segoeui.ttf"]
        )
        candidates.extend(str(folder / name) for name in names)
    elif "darwin" in system:
        candidates.extend(
            [
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
                if bold
                else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Supplemental/Arial.ttf",
            ]
        )
    else:
        # DejaVu is available on Render and supports Arabic + Latin.
        candidates.extend(
            [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
                if bold
                else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/opentype/noto/NotoNaskhArabic-Bold.ttf"
                if bold
                else "/usr/share/fonts/opentype/noto/NotoNaskhArabic-Regular.ttf",
                "/usr/share/fonts/opentype/noto/NotoSansArabic-Bold.ttf"
                if bold
                else "/usr/share/fonts/opentype/noto/NotoSansArabic-Regular.ttf",
                "/usr/share/fonts/opentype/noto/NotoSans-Bold.ttf"
                if bold
                else "/usr/share/fonts/opentype/noto/NotoSans-Regular.ttf",
            ]
        )

    return candidates


def _load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    for candidate in _font_candidates(bold):
        path = Path(candidate)
        if not path.exists():
            continue
        try:
            return ImageFont.truetype(str(path), size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _is_rtl(text: object) -> bool:
    value = str(text or "")
    arabic_count = len(ARABIC_PATTERN.findall(value))
    latin_count = len(LATIN_PATTERN.findall(value))
    return arabic_count > 0 and arabic_count >= latin_count


def _fallback_visual_text(text: str) -> str:
    if not _is_rtl(text) or arabic_reshaper is None or get_display is None:
        return text
    try:
        return get_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


def _layout_args(text: str) -> dict:
    if features.check("raqm"):
        return {
            "direction": "rtl" if _is_rtl(text) else "ltr",
            "language": "ar" if _is_rtl(text) else "en",
        }
    return {}


def _visual_text(text: object) -> str:
    value = str(text or "")
    return value if features.check("raqm") else _fallback_visual_text(value)


def _report_language(analysis: dict) -> str:
    explicit = str(analysis.get("_language", "")).lower()
    if explicit in {"ar", "en"}:
        return explicit
    sample = " ".join(
        str(analysis.get(key, ""))
        for key in ("headline", "summary", "impact", "risks_watchlist")
    )
    arabic_count = len(ARABIC_PATTERN.findall(sample))
    latin_count = len(LATIN_PATTERN.findall(sample))
    return "ar" if arabic_count >= latin_count else "en"


def _analysis_sections(analysis: dict) -> list[tuple[str, str | list[str]]]:
    language = _report_language(analysis)
    labels = (
        {"headline": "عنوان الخبر", "summary": "ملخص الخبر", "takeaways": "أهم النقاط"}
        if language == "ar"
        else {"headline": "News Title", "summary": "News Summary", "takeaways": "Key Points"}
    )
    return [
        (labels["headline"], str(analysis.get("headline", ""))),
        (labels["summary"], str(analysis.get("summary", "") or " ".join(str(item) for item in analysis.get("summary_lines", []) if str(item).strip()))),
        (labels["takeaways"], [str(item) for item in analysis.get("key_takeaways", []) if str(item).strip()][:6]),
    ]


def analysis_text(analysis: dict) -> str:
    title = (
        "تقرير مركز المعلومات بالذكاء الاصطناعي"
        if _report_language(analysis) == "ar"
        else "AI Information Hub Report"
    )
    lines: list[str] = [title, ""]
    for title, content in _analysis_sections(analysis):
        if isinstance(content, list):
            if not content:
                continue
            lines.append(f"{title}:")
            lines.extend(f"- {item}" for item in content)
        else:
            if not content.strip():
                continue
            lines.append(f"{title}: {content}")
        lines.append("")
    return "\n".join(lines).strip()


def _text_bbox(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont):
    visual = _visual_text(text)
    try:
        return draw.textbbox((0, 0), visual, font=font, **_layout_args(text))
    except (TypeError, ValueError):
        return draw.textbbox((0, 0), _fallback_visual_text(text), font=font)


def _text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> int:
    bbox = _text_bbox(draw, text, font)
    return max(0, bbox[2] - bbox[0])


def _wrap_text(
    text: str,
    draw: ImageDraw.ImageDraw,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    lines: list[str] = []
    paragraphs = str(text or "").splitlines() or [""]

    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            lines.append("")
            continue

        words = paragraph.split()
        current: list[str] = []

        for word in words:
            candidate = " ".join(current + [word])
            if _text_width(draw, candidate, font) <= max_width:
                current.append(word)
                continue

            if current:
                lines.append(" ".join(current))
                current = []

            if _text_width(draw, word, font) <= max_width:
                current = [word]
                continue

            piece = ""
            for character in word:
                candidate_piece = piece + character
                if _text_width(draw, candidate_piece, font) <= max_width:
                    piece = candidate_piece
                else:
                    if piece:
                        lines.append(piece)
                    piece = character
            if piece:
                current = [piece]

        if current:
            lines.append(" ".join(current))

    return lines


def _line_height(draw: ImageDraw.ImageDraw, font: ImageFont.FreeTypeFont) -> int:
    bbox = _text_bbox(draw, "أبجدية عربية English 123", font)
    return max(50, (bbox[3] - bbox[1]) + 20)


def _new_page() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), BACKGROUND)
    return image, ImageDraw.Draw(image)


def _draw_text(
    draw: ImageDraw.ImageDraw,
    left: int,
    right: int,
    y: int,
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
    force_rtl: bool | None = None,
) -> None:
    visual = _visual_text(text)
    rtl = _is_rtl(text) if force_rtl is None else force_rtl
    x = right if rtl else left
    anchor = "ra" if rtl else "la"
    try:
        draw.text(
            (x, y),
            visual,
            font=font,
            fill=fill,
            anchor=anchor,
            **({"direction": "rtl", "language": "ar"} if rtl and features.check("raqm") else _layout_args(text)),
        )
    except (TypeError, ValueError):
        fallback = _fallback_visual_text(text)
        if rtl:
            width = _text_width(draw, fallback, font)
            draw.text((right - width, y), fallback, font=font, fill=fill)
        else:
            draw.text((left, y), fallback, font=font, fill=fill)


def _draw_header(draw: ImageDraw.ImageDraw, title_font: ImageFont.FreeTypeFont, language: str) -> int:
    draw.rounded_rectangle(
        (MARGIN, 70, PAGE_WIDTH - MARGIN, 220),
        radius=28,
        fill=HEADER_BACKGROUND,
    )
    title = (
        "تقرير مركز المعلومات بالذكاء الاصطناعي"
        if language == "ar"
        else "AI Information Hub Report"
    )
    try:
        draw.text(
            (PAGE_WIDTH // 2, 145),
            _visual_text(title),
            font=title_font,
            fill=(255, 255, 255),
            anchor="mm",
            **_layout_args(title),
        )
    except (TypeError, ValueError):
        width = _text_width(draw, title, title_font)
        draw.text(
            ((PAGE_WIDTH - width) // 2, 110),
            _fallback_visual_text(title),
            font=title_font,
            fill=(255, 255, 255),
        )
    return 270


def _section_lines(
    content: str | list[str],
    draw: ImageDraw.ImageDraw,
    font: ImageFont.FreeTypeFont,
) -> list[str]:
    if isinstance(content, list):
        output: list[str] = []
        for item in content:
            if str(item).strip():
                output.extend(_wrap_text(f"• {item}", draw, font, CONTENT_WIDTH - 90))
        return output
    if not str(content).strip():
        return []
    return _wrap_text(str(content), draw, font, CONTENT_WIDTH - 90)


def _create_report_pages(analysis: dict) -> list[Image.Image]:
    title_font = _load_font(43, bold=True)
    section_font = _load_font(32, bold=True)
    body_font = _load_font(28, bold=False)
    continuation_font = _load_font(22, bold=True)

    language = _report_language(analysis)
    pages: list[Image.Image] = []
    image, draw = _new_page()
    y = _draw_header(draw, title_font, language)
    body_line_height = _line_height(draw, body_font)
    max_lines_per_card = max(1, (PAGE_HEIGHT - MARGIN - 270 - 165) // body_line_height)

    for section_title, section_content in _analysis_sections(analysis):
        lines = _section_lines(section_content, draw, body_font)
        if not lines:
            continue

        first_chunk = True
        while lines:
            available_height = PAGE_HEIGHT - MARGIN - y
            available_lines = max(1, (available_height - 145) // body_line_height)

            if available_lines < 2:
                pages.append(image)
                image, draw = _new_page()
                y = _draw_header(draw, title_font, language)
                body_line_height = _line_height(draw, body_font)
                available_lines = max_lines_per_card

            chunk_size = min(len(lines), available_lines, max_lines_per_card)
            chunk = lines[:chunk_size]
            lines = lines[chunk_size:]

            card_height = 115 + len(chunk) * body_line_height + 30
            card_bottom = y + card_height

            draw.rounded_rectangle(
                (MARGIN, y, PAGE_WIDTH - MARGIN, card_bottom),
                radius=24,
                fill=CARD_BACKGROUND,
                outline=BORDER_COLOR,
                width=3,
            )

            heading = section_title if first_chunk else f"{section_title} - متابعة"
            _draw_text(
                draw,
                MARGIN + CARD_PADDING,
                PAGE_WIDTH - MARGIN - CARD_PADDING,
                y + 25,
                heading,
                section_font if first_chunk else continuation_font,
                ACCENT_COLOR,
                force_rtl=(language == "ar"),
            )

            text_y = y + 85
            for line in chunk:
                _draw_text(
                    draw,
                    MARGIN + CARD_PADDING,
                    PAGE_WIDTH - MARGIN - CARD_PADDING,
                    text_y,
                    line,
                    body_font,
                    TEXT_COLOR,
                    force_rtl=(language == "ar"),
                )
                text_y += body_line_height

            y = card_bottom + 24
            first_chunk = False

            if lines:
                pages.append(image)
                image, draw = _new_page()
                y = _draw_header(draw, title_font, language)
                body_line_height = _line_height(draw, body_font)

    pages.append(image)
    return pages


def _set_paragraph_direction(paragraph, rtl: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def _set_run_font(run, name: str = "Arial") -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), name)


def _make_docx_uncached(analysis: dict) -> bytes:
    document = Document()
    language = _report_language(analysis)
    report_title = (
        "تقرير مركز المعلومات بالذكاء الاصطناعي"
        if language == "ar"
        else "AI Information Hub Report"
    )
    title = document.add_heading(report_title, level=1)
    _set_paragraph_direction(title, language == "ar")
    for run in title.runs:
        _set_run_font(run)

    for section_title, section_content in _analysis_sections(analysis):
        if isinstance(section_content, list) and not section_content:
            continue
        if isinstance(section_content, str) and not section_content.strip():
            continue

        heading = document.add_heading(section_title, level=2)
        _set_paragraph_direction(heading, language == "ar")
        for run in heading.runs:
            _set_run_font(run)

        values = section_content if isinstance(section_content, list) else [section_content]
        for item in values:
            text = str(item)
            paragraph = document.add_paragraph(style="List Bullet" if isinstance(section_content, list) else None)
            run = paragraph.add_run(text)
            _set_run_font(run)
            _set_paragraph_direction(paragraph, _is_rtl(text))

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _make_pdf_uncached(analysis: dict) -> bytes:
    pages = _create_report_pages(analysis)
    output = BytesIO()
    first = pages[0].convert("RGB")
    remaining = [page.convert("RGB") for page in pages[1:]]
    first.save(
        output,
        format="PDF",
        resolution=150.0,
        save_all=True,
        append_images=remaining,
    )
    return output.getvalue()


def _make_image_uncached(analysis: dict) -> bytes:
    pages = _create_report_pages(analysis)
    gap = 36
    total_height = sum(page.height for page in pages) + gap * (len(pages) - 1)
    canvas = Image.new("RGB", (PAGE_WIDTH, total_height), BACKGROUND)
    y = 0
    for page in pages:
        canvas.paste(page, (0, y))
        y += page.height + gap
    output = BytesIO()
    canvas.save(output, format="PNG", optimize=True)
    return output.getvalue()


def _canonical_payload(analysis: dict) -> str:
    return json.dumps(analysis, ensure_ascii=False, sort_keys=True, default=str)


@lru_cache(maxsize=64)
def _cached_pdf(payload: str) -> bytes:
    return _make_pdf_uncached(json.loads(payload))


@lru_cache(maxsize=64)
def _cached_docx(payload: str) -> bytes:
    return _make_docx_uncached(json.loads(payload))


@lru_cache(maxsize=64)
def _cached_image(payload: str) -> bytes:
    return _make_image_uncached(json.loads(payload))


def make_pdf(analysis: dict) -> bytes:
    return _cached_pdf(_canonical_payload(analysis))


def make_docx(analysis: dict) -> bytes:
    return _cached_docx(_canonical_payload(analysis))


def make_image(analysis: dict) -> bytes:
    return _cached_image(_canonical_payload(analysis))


@lru_cache(maxsize=64)
def _cached_text(payload: str) -> bytes:
    return analysis_text(json.loads(payload)).encode("utf-8-sig")


@lru_cache(maxsize=64)
def _cached_json(payload: str) -> bytes:
    pretty = json.dumps(json.loads(payload), ensure_ascii=False, indent=2)
    return pretty.encode("utf-8-sig")


def make_text(analysis: dict) -> bytes:
    return _cached_text(_canonical_payload(analysis))


def make_json(analysis: dict) -> bytes:
    return _cached_json(_canonical_payload(analysis))
